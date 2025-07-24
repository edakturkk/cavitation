from docx import Document
import os

def read_docx_detailed(file_path):
    """Word dosyasını detaylı olarak oku"""
    try:
        doc = Document(file_path)
        content = []
        
        print(f"Dosya açıldı: {file_path}")
        print(f"Toplam paragraf sayısı: {len(doc.paragraphs)}")
        print(f"Toplam tablo sayısı: {len(doc.tables)}")
        print(f"Toplam bölüm sayısı: {len(doc.sections)}")
        print("-" * 50)
        
        # Paragrafları oku
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:  # Boş olmayan paragrafları al
                content.append(f"Paragraf {i+1}: {text}")
                print(f"Paragraf {i+1}: {text[:100]}...")  # İlk 100 karakteri göster
        
        # Tabloları oku
        for i, table in enumerate(doc.tables):
            content.append(f"\n--- TABLO {i+1} ---")
            print(f"\n--- TABLO {i+1} ---")
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                row_content = " | ".join(row_text)
                content.append(row_content)
                print(row_content)
        
        return '\n'.join(content)
    
    except Exception as e:
        return f"Dosya okuma hatası: {str(e)}"

# Dosya yolunu belirt
file_path = "dosyalar/EDA EN SON HALİ FORM-31_TEZ KONUSU ve PLANI (2) (3).docx"

# Dosyanın var olup olmadığını kontrol et
if os.path.exists(file_path):
    print("Dosya bulundu, detaylı okuma başlıyor...")
    content = read_docx_detailed(file_path)
    
    # Tam içeriği dosyaya kaydet
    with open("dosya_icerigi.txt", "w", encoding="utf-8") as f:
        f.write(content)
    
    print(f"\nTam içerik 'dosya_icerigi.txt' dosyasına kaydedildi.")
    
else:
    print(f"Dosya bulunamadı: {file_path}") 